import re
from docx import Document
from typing import List, Dict

def extract_variables(file_path: str) -> Dict[str, List[str]]:
    """
    Извлекает переменные из документа Word.
    Переменные должны быть в формате ##имя_переменной##
    
    Args:
        file_path: путь к файлу документа
        
    Returns:
        Dict с ключами:
        - variables: список найденных переменных
        - is_template: является ли документ шаблоном
    """
    try:
        doc = Document(file_path)
        variables = set()
        
        # Ищем переменные во всех параграфах
        for paragraph in doc.paragraphs:
            matches = re.findall(r'##([^#]+)##', paragraph.text)
            variables.update(matches)
            
        # Ищем переменные в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'##([^#]+)##', cell.text)
                    variables.update(matches)
        
        return {
            "variables": list(variables),
            "is_template": len(variables) > 0
        }
    except Exception as e:
        return {
            "variables": [],
            "is_template": False,
            "error": str(e)
        }

def extract_text_from_docx(file_path):
    """
    Извлекает весь текст из файла docx
    """
    doc = Document(file_path)
    full_text = []
    
    # Извлекаем текст из параграфов
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text) 