from typing import List, Any, Dict
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import os
from datetime import datetime
from docx import Document
import re
from pydantic import BaseModel

from app import crud
from app.db.database import get_db
from app.schemas.template import Template, TemplateCreate
from app.api import deps
from app.models.user import User
from app.services.google_docs import GoogleDocsService
from app.utils.document_parser import extract_text_from_docx

router = APIRouter()

# В начале файла добавим создание директории, если её нет
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class GoogleDocsRequest(BaseModel):
    variables: Dict[str, str]

class TemplateContentUpdate(BaseModel):
    content: str

@router.post("/", response_model=Template)
async def create_template(
    *,
    db: Session = Depends(deps.get_db),
    file: UploadFile = File(...),
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Загрузить новый шаблон.
    """
    # Сохраняем файл
    file_path = os.path.join(UPLOAD_DIR, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}")
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Создаем запись в базе данных
    template_in = TemplateCreate(
        filename=file.filename,
        file_path=file_path,
        content_type=file.content_type,
        user_id=current_user.id,
        is_template=False,
        variables=[]
    )
    
    return crud.template.create_with_variables(db=db, obj_in=template_in, file_path=file_path)

@router.get("/", response_model=List[Template])
def read_templates(
    db: Session = Depends(get_db),
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Получить список шаблонов пользователя.
    """
    templates = crud.template.get_by_user(db, user_id=current_user.id, skip=skip, limit=limit)
    return templates

@router.get("/{template_id}", response_model=Template)
def read_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_user)
):
    template = crud.template.get(db, id=template_id)
    if template is None or template.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Template not found")
    return template

@router.delete("/{template_id}", response_model=Template)
def delete_template(
    template_id: int,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Удалить шаблон.
    """
    template = crud.template.get(db, id=template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    if template.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    return crud.template.remove(db=db, id=template_id)

@router.get("/{template_id}/download")
def download_template(
    template_id: int,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Скачать шаблон.
    """
    template = crud.template.get(db, id=template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    if template.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    
    if not os.path.exists(template.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        template.file_path,
        filename=template.filename,
        media_type=template.content_type
    )

@router.post("/{template_id}/generate")
async def generate_document(
    template_id: int,
    variables: Dict[str, str],
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Сгенерировать заполненный документ из шаблона.
    """
    try:
        template = crud.template.get(db, id=template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        if template.user_id != current_user.id:
            raise HTTPException(status_code=403, detail="Not enough permissions")
        
        # Проверяем существование файла
        if not os.path.exists(template.file_path):
            error_msg = f"Template file not found: {template.file_path}"
            print(error_msg)
            raise HTTPException(status_code=404, detail=error_msg)
        
        # Создаем директорию для выходных файлов, если она не существует
        output_dir = "uploads/generated"
        os.makedirs(output_dir, exist_ok=True)
        output_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{template.filename}"
        output_path = os.path.join(output_dir, output_filename)
        
        # Обработка в зависимости от типа файла
        if template.file_path.endswith('.docx'):
            try:
                # Создаем копию файла
                import shutil
                shutil.copy2(template.file_path, output_path)
                
                # Открываем копию для редактирования
                doc = Document(output_path)
                
                # Заменяем все переменные на значения в параграфах
                for paragraph in doc.paragraphs:
                    original_text = paragraph.text
                    modified_text = original_text
                    
                    # Ищем все переменные и заменяем их
                    for var_name in template.variables:
                        placeholder = f"##{var_name}##"
                        if placeholder in modified_text:
                            modified_text = modified_text.replace(
                                placeholder, 
                                variables.get(var_name, "")
                            )
                    
                    # Если текст изменился, обновляем параграф
                    if modified_text != original_text:
                        # Очищаем все runs в параграфе и добавляем новый
                        for i in range(len(paragraph.runs)):
                            paragraph.runs[0].text = ""
                        
                        paragraph.add_run(modified_text)
                
                # Заменяем переменные в таблицах
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                original_text = paragraph.text
                                modified_text = original_text
                                
                                # Ищем все переменные и заменяем их
                                for var_name in template.variables:
                                    placeholder = f"##{var_name}##"
                                    if placeholder in modified_text:
                                        modified_text = modified_text.replace(
                                            placeholder, 
                                            variables.get(var_name, "")
                                        )
                                
                                # Если текст изменился, обновляем параграф
                                if modified_text != original_text:
                                    # Очищаем все runs в параграфе и добавляем новый
                                    for i in range(len(paragraph.runs)):
                                        paragraph.runs[0].text = ""
                                    
                                    paragraph.add_run(modified_text)
                
                # Сохраняем результат
                doc.save(output_path)
                
                return FileResponse(
                    output_path,
                    filename=output_filename,
                    media_type=template.content_type
                )
            except Exception as e:
                error_msg = f"Error processing docx: {str(e)}"
                print(error_msg)
                raise HTTPException(status_code=500, detail=error_msg)
        else:
            # Для текстовых файлов
            try:
                # Читаем содержимое
                with open(template.file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Заменяем переменные
                for var_name in template.variables:
                    placeholder = f"##{var_name}##"
                    content = content.replace(placeholder, variables.get(var_name, ""))
                
                # Сохраняем результат
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                return FileResponse(
                    output_path,
                    filename=output_filename,
                    media_type=template.content_type
                )
            except Exception as e:
                error_msg = f"Error processing text file: {str(e)}"
                print(error_msg)
                raise HTTPException(status_code=500, detail=error_msg)
    except HTTPException:
        raise
    except Exception as e:
        error_msg = f"Failed to generate document: {str(e)}"
        print(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

@router.post("/{template_id}/google-docs")
async def create_google_doc(
    template_id: int,
    variables: Dict[str, str],
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
):
    try:
        # Получаем шаблон
        template = crud.template.get(db, id=template_id)
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        if template.user_id != current_user.id:
            raise HTTPException(status_code=403, detail="Not enough permissions")

        # Проверяем токен Google
        if not current_user.google_token:
            raise HTTPException(
                status_code=401,
                detail="Google authentication required"
            )

        # Читаем содержимое файла шаблона
        content = None
        if template.file_path.endswith('.docx'):
            content = extract_text_from_docx(template.file_path)
        else:
            try:
                # Пробуем разные кодировки
                encodings = ['utf-8', 'cp1251', 'latin-1']
                
                for encoding in encodings:
                    try:
                        with open(template.file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        break  # Если успешно прочитано, выходим из цикла
                    except UnicodeDecodeError:
                        continue
                    
                if content is None:
                    # Если не удалось прочитать ни с одной кодировкой, читаем в бинарном режиме
                    with open(template.file_path, 'rb') as f:
                        content = f.read().decode('utf-8', errors='replace')
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error reading file: {str(e)}")

        # Заменяем переменные в контенте
        for key, value in variables.items():
            content = content.replace(f"##{key}##", value)

        # Создаем Google Doc
        google_service = GoogleDocsService(
            access_token=current_user.google_token,
            refresh_token=current_user.google_refresh_token
        )
        result = await google_service.create_document(
            title=f"{template.filename} - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            content=content
        )

        # Логируем результат
        print("Created document:", result)

        return result
    except Exception as e:
        print(f"Error in create_google_doc: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create Google Doc: {str(e)}"
        )

@router.get("/{template_id}/content")
async def get_template_content(
    template_id: int,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
):
    """
    Получить содержимое шаблона.
    """
    # Проверяем существование шаблона и права доступа
    template = crud.template.get(db, id=template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    if template.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    
    # Проверка на тип файла
    is_binary = template.content_type not in [
        "text/plain", "text/html", "text/markdown", "application/json", 
        "text/css", "text/javascript", "text/csv"
    ]
    
    # Проверяем существование файла
    if not os.path.exists(template.file_path):
        error_msg = f"File not found: {template.file_path}"
        print(error_msg)
        raise HTTPException(status_code=404, detail=error_msg)
    
    # Для бинарных форматов (DOCX, PDF, и т.д.) извлекаем текст и переменные
    if is_binary or template.file_path.endswith(('.docx', '.pdf')):
        # Извлекаем текст из документа Word
        if template.file_path.endswith('.docx'):
            try:
                # Извлекаем полный текст из docx
                doc = Document(template.file_path)
                
                # Извлекаем текст из всех параграфов, сохраняя структуру
                full_text = ""
                for para in doc.paragraphs:
                    if para.text.strip():  # Если параграф не пустой
                        full_text += para.text + "\n\n"
                
                # Извлекаем текст из таблиц
                for table in doc.tables:
                    full_text += "\n--- Table ---\n"
                    for row in table.rows:
                        row_text = ""
                        for cell in row.cells:
                            row_text += cell.text + " | "
                        full_text += row_text[:-3] + "\n"  # Убираем последний разделитель " | "
                    full_text += "--- End Table ---\n\n"
                
                # Добавляем информацию о переменных
                variables = template.variables if template.variables else []
                
                # Добавляем инструкцию
                instructions = (
                    "This is a preview of the Word document content. "
                    "When you save, a new text template will be created "
                    "that you can fully edit.\n\n"
                    "You can add or modify variables using ##variable_name## syntax.\n\n"
                    "--- DOCUMENT CONTENT ---\n\n"
                )
                
                content = instructions + full_text
                
                # Если в тексте нет переменных, добавляем примеры
                if not variables:
                    content += "\n\n--- EXAMPLE VARIABLES ---\n"
                    content += "You can add variables like ##client_name## or ##date## anywhere in the text.\n"
                
                return {"content": content, "is_binary": True}
            except Exception as e:
                print(f"Error extracting text from docx: {str(e)}")
                return {
                    "content": (
                        "Could not extract content from Word document. "
                        "When you save, a new text template will be created "
                        "that you can fully edit with variables like ##variable_name##."
                    ), 
                    "is_binary": True
                }
        else:
            return {"content": "This file type cannot be edited as text.", "is_binary": True}
    
    # Для текстовых файлов читаем содержимое
    try:
        content = None
        # Пробуем разные кодировки
        encodings = ['utf-8', 'cp1251', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(template.file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break  # Если успешно прочитано, выходим из цикла
            except UnicodeDecodeError:
                continue
            
        if content is None:
            # Если не удалось прочитать ни с одной кодировкой, обрабатываем как бинарный файл
            return {"content": "This file cannot be read as text. It may be a binary file.", "is_binary": True}
                
        return {"content": content, "is_binary": False}
    except Exception as e:
        error_msg = f"Error reading template content: {str(e)}"
        print(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

@router.put("/{template_id}/content")
async def update_template_content(
    template_id: int,
    content_update: TemplateContentUpdate,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user)
):
    """
    Обновить содержимое шаблона.
    """
    # Проверяем существование шаблона и права доступа
    template = crud.template.get(db, id=template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    if template.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    
    # Проверка на тип файла
    is_binary = template.content_type not in [
        "text/plain", "text/html", "text/markdown", "application/json", 
        "text/css", "text/javascript", "text/csv"
    ]
    
    # Если это бинарный файл (docx), создаем новый текстовый шаблон
    if is_binary or template.file_path.endswith(('.docx', '.pdf')):
        try:
            # Если пользователь уже отредактировал содержимое, используем его
            # Иначе, извлекаем текст из оригинального документа
            template_content = content_update.content
            
            # Если в содержимом есть только список переменных (что видно по предупреждению),
            # то извлекаем полный текст из оригинального документа
            if "This is a Word document and cannot be edited directly as text" in template_content:
                if template.file_path.endswith('.docx'):
                    try:
                        # Извлекаем полный текст из docx
                        doc = Document(template.file_path)
                        
                        # Извлекаем текст из всех параграфов, сохраняя структуру
                        full_text = ""
                        for para in doc.paragraphs:
                            if para.text.strip():  # Если параграф не пустой
                                full_text += para.text + "\n\n"
                        
                        # Извлекаем текст из таблиц
                        for table in doc.tables:
                            full_text += "\n--- Table ---\n"
                            for row in table.rows:
                                row_text = ""
                                for cell in row.cells:
                                    row_text += cell.text + " | "
                                full_text += row_text[:-3] + "\n"  # Убираем последний разделитель " | "
                            full_text += "--- End Table ---\n\n"
                        
                        # Используем этот текст как основу для нового шаблона
                        template_content = full_text
                    except Exception as e:
                        print(f"Error extracting text from docx: {str(e)}")
                        # Если не удалось извлечь текст, оставляем содержимое как есть
            
            # Извлекаем переменные из содержимого
            variables = []
            import re
            for match in re.finditer(r'##([^#]+)##', template_content):
                var_name = match.group(1)
                if var_name not in variables:
                    variables.append(var_name)
            
            # Создаем новый файл
            new_filename = f"{os.path.splitext(template.filename)[0]}_text.txt"
            new_file_path = os.path.join(UPLOAD_DIR, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{new_filename}")
            
            # Записываем содержимое в новый файл
            with open(new_file_path, 'w', encoding='utf-8') as f:
                f.write(template_content)
            
            # Создаем запись в базе данных для нового шаблона
            template_in = {
                "filename": new_filename,
                "file_path": new_file_path,
                "content_type": "text/plain",
                "user_id": current_user.id,
                "is_template": len(variables) > 0,
                "variables": variables
            }
            
            # Создаем новый шаблон
            new_template = crud.template.create(db=db, obj_in=template_in)
            
            # Возвращаем информацию о новом шаблоне
            return {
                "message": "Created new text template from binary file",
                "original_template": template,
                "new_template": new_template
            }
        except Exception as e:
            error_msg = f"Error creating new template: {str(e)}"
            print(error_msg)
            raise HTTPException(status_code=500, detail=error_msg)
    
    # Для текстовых файлов обновляем содержимое файла
    try:
        with open(template.file_path, 'w', encoding='utf-8') as f:
            f.write(content_update.content)
        
        # Извлекаем переменные из содержимого
        variables = []
        import re
        for match in re.finditer(r'##([^#]+)##', content_update.content):
            var_name = match.group(1)
            if var_name not in variables:
                variables.append(var_name)
        
        # Обновляем запись в базе данных
        template_data = {"variables": variables, "is_template": len(variables) > 0}
        updated_template = crud.template.update(db, db_obj=template, obj_in=template_data)
        
        return updated_template
    except Exception as e:
        error_msg = f"Error updating template content: {str(e)}"
        print(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)
